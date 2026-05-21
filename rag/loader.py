from pypdf import PdfReader
from docx import Document
import io
import os

MAX_FILE_SIZE = 10*1024*1024

def load(file_bytes: bytes , filename : str) -> str:


    filename = os.path.basename(filename)
    if not filename:
        raise ValueError("Invalid file name")

    if len(file_bytes)==0:
        raise ValueError("Empty file")

    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError("File size exceeds 10MB")


    if filename.endswith(".pdf"):
        text = load_pdf(file_bytes)

    elif filename.endswith(".docx"):
        text = load_docx(file_bytes)

    elif filename.endswith(".txt"):
        text = load_txt(file_bytes)

    else:
        raise ValueError(f"unsupported fileType : {filename}")

    if not text.strip():
        raise ValueError(f"No extractable text found in {filename} , file may be image based empty")

    return text


def load_pdf(file_bytes:bytes)-> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    text=""
    for page in reader.pages:
        text += page.extract_text() or ""

    return text


def load_docx(file_bytes : bytes)-> str:
    doc = Document(io.BytesIO(file_bytes))
    text=""
    for para in doc.paragraphs:
        text += para.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text +="\n"

    return text

def load_txt(file_bytes : bytes)-> str:
    try :
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")




